#!/usr/bin/env python3
"""Inspect headings and nearby paragraphs in a DOCX file."""

import argparse
from pathlib import Path

from docx import Document


def iter_blocks(doc):
    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue
        style = paragraph.style.name if paragraph.style is not None else ""
        yield idx, style, text


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("docx_path")
    parser.add_argument("--window", type=int, default=12)
    args = parser.parse_args()

    path = Path(args.docx_path)
    doc = Document(str(path))

    print(f"DOCX: {path}")
    print(f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} sections={len(doc.sections)}")
    print("\nHEADINGS")
    for idx, style, text in iter_blocks(doc):
        if style.startswith("Heading") or text[:4].replace(".", "").isdigit():
            print(f"{idx:04d}\t{style}\t{text[:180]}")

    targets = []
    for idx, style, text in iter_blocks(doc):
        stripped = text.replace(" ", "")
        if stripped.startswith("3.3") or stripped.startswith("4"):
            targets.append(idx)

    print("\nTARGET WINDOWS")
    seen = set()
    for target in targets:
        start = max(0, target - args.window)
        end = min(len(doc.paragraphs), target + args.window + 1)
        for i in range(start, end):
            if i in seen:
                continue
            seen.add(i)
            p = doc.paragraphs[i]
            text = p.text.strip()
            if text:
                style = p.style.name if p.style is not None else ""
                print(f"{i:04d}\t{style}\t{text}")


if __name__ == "__main__":
    main()
