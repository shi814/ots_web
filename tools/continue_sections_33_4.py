#!/usr/bin/env python3
"""Insert drafted text for sections 3.3 and 4 into the target DOCX."""

import argparse
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


SECTION_33_PARAGRAPHS = [
    (
        "综合上述实验结果，本文方法的主要优势首先体现在 OTS 元件约束与神经网络输出空间的一致性。"
        "传统连续参数生成方法通常需要在网络输出后再进行最近邻匹配或库内替换，这一过程容易改变原有像差平衡，"
        "也可能使连续预测结果与真实可采购元件之间产生偏差。相比之下，本文将玻璃面直接建模为 OTS 候选组内的分类选择，"
        "使曲率、厚度、玻璃材料及折射率色散等参数在生成阶段即来自真实库内元件。"
        "因此，模型输出不只是数值上接近某个目录元件，而是在结构上天然满足库内合法性和可采购约束。"
    ),
    (
        "其次，本文将离散镜片选择与连续空气层精修分解为两个相互衔接的阶段，降低了 OTS 混合优化问题的难度。"
        "第一阶段侧重于在库内候选空间中生成具备基本成像能力的初始结构，第二阶段则在固定玻璃处方的基础上仅调整空气间隔。"
        "这种处理方式保留了 OTS 元件选择的离散稳定性，同时利用空气层这一连续自由度进一步改善焦距匹配和像差分配。"
        "实验结果表明，经过空气层二次无监督训练后，EFL 筛选后的有效系统数量由 666 个增加到 703 个，"
        "RMS spot radius 均值由约 10.45 μm 降低到 8.26 μm，中位数由 8.58 μm 降低到 6.99 μm，"
        "说明该二阶段策略能够在不改变 OTS 玻璃元件选择的情况下进一步提升系统聚焦性能。"
    ),
    (
        "第三，本文的训练过程不依赖逐系统的人工最优处方标注，而是通过可微光线追迹构造物理自监督信号。"
        "RMS 点列、有效焦距误差、畸变、远心性、光线有效性和表面重叠等指标共同参与损失计算，"
        "使模型在生成过程中同时受到成像质量、一阶规格和物理可行性的约束。"
        "这使该方法区别于仅学习已有设计参数分布的监督回归，也区别于完全依赖穷举或启发式搜索的 OTS 组合优化。"
        "从训练曲线和 Zemax 验证样例可以看出，可微物理反馈能够为空气层预测提供稳定的优化方向，"
        "并在具体光学系统中体现为较低的 RMS 点列和合理的畸变控制。"
    ),
    (
        "尽管如此，当前方法的误差来源仍然较为明确。首先，OTS 元件库本身是离散且有限的，"
        "某些 F-number、HFOV 或表面数量组合可能缺少足够合适的库内曲率、厚度和玻璃材料搭配，"
        "这会限制模型能够达到的最优像差平衡。即使空气间隔可以连续调整，当玻璃元件组合已经固定时，"
        "二次训练也只能在该固定处方附近重新分配光焦度和像差，无法像完整连续设计那样自由改变每个玻璃面的曲率或材料。"
        "因此，部分系统在畸变和远心性上仍可能接近或超过阈值，这并非单纯的训练不充分，而是 OTS 候选空间离散性的直接体现。"
    ),
    (
        "其次，可微物理评价与 Zemax 等商业光学软件中的完整分析之间仍存在建模差异。"
        "本文训练阶段采用有限视场、有限光瞳采样和固定评价指标来构造损失，主要关注 RMS 点列、EFL、畸变和远心性等关键约束；"
        "而实际工程设计还可能涉及更密集的视场采样、更多波长、像面位置优化、光阑机械限制、公差、杂散光、镀膜和装调误差等因素。"
        "因此，网络输出通过可微追迹获得较好指标并不意味着可以完全替代 Zemax 中的最终验证与工程优化。"
        "本文中的 Zemax 对比主要用于验证生成结构的物理可行性和像质趋势，而不是给出完整量产级设计结论。"
    ),
    (
        "此外，本文采用的硬筛选策略也会影响统计结果的解释。EFL loss 筛选能够保证进入统计分析的系统满足基本焦距要求，"
        "使 RMS、畸变和远心性等指标比较更加有意义；但筛选后的样本分布并不等同于模型对所有输入规格的无条件输出分布。"
        "因此，后续工作可进一步分析未通过筛选样本的失败模式，并将筛选结果回流到模型训练或候选库构建中，"
        "形成更主动的样本挖掘机制。与此同时，当前实验主要围绕扫描镜头展开，系统规格、材料序列和表面数量仍相对有限，"
        "方法在更大规模 OTS 库、更复杂多组元结构以及其他类型光学系统中的泛化能力仍需要进一步验证。"
    ),
]


SECTION_4_PARAGRAPHS = [
    (
        "本文围绕 OTS 光学元件库约束下的扫描镜头自动设计问题，提出了一种 library-native 的无监督 Transformer 分类生成框架。"
        "与传统连续参数回归后再进行元件匹配的方法不同，本文将玻璃表面直接建模为 OTS 候选组内的分类选择，"
        "使网络输出的曲率、厚度、玻璃材料和折射率色散等参数在生成阶段即来自真实可采购元件库；"
        "同时，将空气层建模为连续回归变量，并通过可微光线追迹构建 RMS 点列、有效焦距误差、畸变、远心性、"
        "光线有效性和表面重叠惩罚等无监督物理损失。该框架使离散库内元件选择与连续空气层优化能够在统一的物理反馈链路中协同工作。"
    ),
    (
        "在第一阶段 OTS 约束初始生成实验中，模型能够直接输出由库内真实元件组成的扫描镜头结构。"
        "经过有效焦距约束筛选后，共得到 666 个生成系统；这些系统在 50 μm RMS 阈值下全部满足要求，"
        "畸变和像方远心性也保持较高合格率，说明所提出的分类生成模型能够在离散 OTS 候选空间内学习到具有基本成像能力的结构规律。"
        "代表性 Zemax 对比进一步表明，生成系统不仅在统计指标上满足阈值，也能够形成合理光路布局和有效点列分布，"
        "具备作为 OTS 初始设计的可用性。"
    ),
    (
        "在第二阶段空气层二次无监督训练中，本文固定第一阶段得到的 OTS 玻璃元件选择及其库内处方，"
        "仅更新空气层回归分支，并通过可微光线追迹重新计算物理损失。实验结果表明，二次训练后的损失曲线逐渐下降并趋于稳定，"
        "说明可微物理评价能够为空气层连续参数提供有效优化反馈。与第一阶段相比，EFL 筛选后的有效系统数量由 666 个增加到 703 个，"
        "RMS spot radius 分布进一步向小半径区域集中，均值和中位数均明显降低。"
        "Zemax 代表性样例也显示，二次训练后的生成系统在中心视场、中间视场和边缘视场下均获得更小的 RMS 点列，"
        "验证了固定玻璃处方下空气层物理自监督精修的有效性。"
    ),
    (
        "总体而言，本文方法为 OTS 光学系统设计提供了一种区别于穷举搜索、启发式优化和连续参数后匹配的新思路。"
        "其核心价值在于将真实元件库作为网络的合法输出空间，而不是作为后处理约束；同时利用可微光线追迹将物理评价嵌入训练过程，"
        "在缺少大规模人工最优标签的情况下实现库内候选筛选和空气层连续参数优化。"
        "该框架能够在保持可采购性和结构合法性的前提下快速生成具有成像能力的扫描镜头候选系统，"
        "并为后续 Zemax 精修、工程约束加入和实验验证提供较好的初始结构。"
    ),
    (
        "未来工作可从三个方向进一步扩展。第一，可引入更大规模和更多厂商的 OTS 元件库，"
        "并将机械口径、外径、边厚、装配间隙和库存可获得性等工程约束纳入候选组构建。"
        "第二，可进一步增强可微物理评价模型，在训练阶段加入更密集的视场和波长采样、公差敏感性、像面位置调整以及多目标权重自适应策略。"
        "第三，可将本文框架与传统局部优化器、Zemax 批量验证或主动学习机制结合，使未通过筛选的系统反向指导候选库重构和网络再训练。"
        "通过这些扩展，library-native 的 OTS 神经网络设计方法有望从扫描镜头推广到显微物镜、机器视觉镜头、光谱仪和其他多元件工程光学系统。"
    ),
]


def insert_paragraph_after(paragraph, text):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    inserted = Paragraph(new_p, paragraph._parent)
    inserted.style = paragraph.style
    inserted.add_run(text)
    return inserted


def find_paragraph(doc, prefix):
    compact_prefix = prefix.replace(" ", "")
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().replace(" ", "").startswith(compact_prefix):
            return paragraph
    raise ValueError(f"Could not find paragraph starting with {prefix!r}")


def insert_section_after_heading(doc, heading_prefix, paragraphs):
    cursor = find_paragraph(doc, heading_prefix)
    for text in reversed(paragraphs):
        insert_paragraph_after(cursor, text)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("input_docx")
    parser.add_argument("output_docx")
    args = parser.parse_args()

    doc = Document(args.input_docx)
    insert_section_after_heading(doc, "4 总结", SECTION_4_PARAGRAPHS)
    insert_section_after_heading(doc, "3.3 方法优势、误差来源与局限性", SECTION_33_PARAGRAPHS)

    output_path = Path(args.output_docx)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"Saved {output_path}")


if __name__ == "__main__":
    main()
